"""Convert a Markdown report to a .docx file (headings, tables, lists, code,
inline bold/italic/code). Purpose-built for the reports in this repo — not a
general Markdown engine, but it covers everything they use.

Usage:  python docs/md_to_docx.py docs/PROJECT_REPORT.md docs/PROJECT_REPORT.docx
"""

from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def add_inline(paragraph, text: str) -> None:
    """Render inline **bold**, *italic*, and `code` spans into ``paragraph``."""
    token = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")
    for part in token.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.color.rgb = RGBColor(0xB0, 0x2A, 0x37)
        else:
            paragraph.add_run(part)


def emit_table(doc: Document, rows: list[str]) -> None:
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    header, body = cells[0], cells[2:]  # row 1 is the |---| separator
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        add_inline(table.rows[0].cells[i].paragraphs[0], h)
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
    for row in body:
        rc = table.add_row().cells
        for i, val in enumerate(row[:len(header)]):
            add_inline(rc[i].paragraphs[0], val)
    doc.add_paragraph()


def convert(md_path: str, docx_path: str) -> None:
    lines = open(md_path, encoding="utf-8").read().splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code blocks.
        if line.strip().startswith("```"):
            i += 1
            code = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            para = doc.add_paragraph()
            run = para.add_run("\n".join(code))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            i += 1
            continue

        # Tables (a line starting with | followed by a |---| separator).
        if line.strip().startswith("|") and i + 1 < len(lines) and re.match(
                r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            emit_table(doc, block)
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue
        if stripped == "---":
            i += 1
            continue

        # Headings.
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            heading = doc.add_heading(level=min(level, 4))
            add_inline(heading, m.group(2))
            i += 1
            continue

        # Bullet / numbered lists.
        if re.match(r"^[-*]\s+", stripped):
            add_inline(doc.add_paragraph(style="List Bullet"), stripped[2:])
            i += 1
            continue
        num = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if num:
            add_inline(doc.add_paragraph(style="List Number"), num.group(2))
            i += 1
            continue

        # Plain paragraph.
        add_inline(doc.add_paragraph(), stripped)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "docs/PROJECT_REPORT.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else "docs/PROJECT_REPORT.docx"
    convert(src, dst)
